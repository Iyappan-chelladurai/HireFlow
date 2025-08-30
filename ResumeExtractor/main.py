from fastapi import FastAPI, File, Form, UploadFile
import io
import os
import re
import docx
import docx2txt
import PyPDF2
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import spacy
from fuzzywuzzy import fuzz
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Load SpaCy model
nlp = spacy.load("en_core_web_sm")

app = FastAPI()

# ---------- File Readers ----------
def extract_text_from_pdf(file_stream):
    """Extract text from PDF with OCR fallback"""
    text = ""
    pdf_reader = PyPDF2.PdfReader(file_stream)

    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    # If no text (scanned PDF), use OCR
    if not text.strip():
        file_stream.seek(0)
        images = convert_from_bytes(file_stream.read())
        for img in images:
            text += pytesseract.image_to_string(img) + "\n"

    return text.strip()


def extract_text_from_docx(file_stream):
    """Extract text from DOCX (paragraphs + tables), fallback to docx2txt"""
    file_stream.seek(0)
    try:
        doc = docx.Document(file_stream)
        full_text = []

        # Normal paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Text inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        full_text.append(cell_text)

        text = "\n".join(full_text).strip()
        if text:
            return text
    except Exception:
        pass

    # Fallback to docx2txt
    file_stream.seek(0)
    return docx2txt.process(file_stream) or ""


# ---------- Data Extractors ----------
def extract_contact_info(text):
    """Extract name, email, phone"""
    email = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    phone = re.search(r"\+?\d[\d\-\(\) ]{8,}\d", text)

    name = None
    doc = nlp(text[:400])  # limit to first 400 chars
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            break

    # fallback: first clean line
    if not name:
        for line in text.split("\n"):
            if line.strip() and len(line.split()) <= 4 and "@" not in line and not line.strip().isdigit():
                name = line.strip()
                break

    return {
        "name": name,
        "email": email.group(0) if email else None,
        "phone": phone.group(0) if phone else None
    }


def extract_skills(text, skills_list):
    """Match skills using fuzzy matching"""
    found_skills = []
    for skill in skills_list:
        if fuzz.token_set_ratio(skill.lower(), text.lower()) > 85:
            found_skills.append(skill)
    return list(set(found_skills))


def extract_education(text):
    """Extract education lines"""
    edu_pattern = r"(B\.?Tech|M\.?Tech|B\.?Sc|M\.?Sc|Bachelor|Master|Ph\.?D|Diploma)"
    return [line.strip() for line in text.split("\n") if re.search(edu_pattern, line, re.IGNORECASE)]


def extract_experience(text):
    """Extract experience lines (avoid project desc noise)"""
    exp_pattern = r"(Intern|Engineer|Developer|Manager|Consultant|Designer|Lead|Associate)"
    results = []
    for line in text.split("\n"):
        if re.search(exp_pattern, line, re.IGNORECASE):
            if "Project" not in line and "Desc" not in line:
                results.append(line.strip())
    return results


def extract_certifications(text):
    """Extract certification lines (filter out Microsoft Office)"""
    cert_keywords = ["certified", "certification", "certificate", "AWS", "Azure", "Google Cloud", "Microsoft"]
    results = []
    for line in text.split("\n"):
        clean_line = line.strip()
        if any(k.lower() in clean_line.lower() for k in cert_keywords):
            if "office" not in clean_line.lower():  # avoid false positives
                results.append(clean_line)
    return results


# ---------- Ranking & Matching ----------
def calculate_job_match_score(resume_text, job_description):
    vectorizer = TfidfVectorizer().fit_transform([resume_text, job_description])
    similarity = cosine_similarity(vectorizer[0:1], vectorizer[1:2])
    return round(float(similarity[0][0]) * 100, 2)


def calculate_skill_match_score(found_skills, required_skills):
    if not required_skills:
        return 0
    match_count = len(set(found_skills) & set(required_skills))
    return round((match_count / len(required_skills)) * 100, 2)


# ---------- Main Parser ----------
def parse_resume(file_stream, filename, skills_list, job_description):
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(file_stream)
    elif ext == ".docx":
        text = extract_text_from_docx(file_stream)
    else:
        raise ValueError("Unsupported file format")

    if not text.strip():
        raise ValueError("No text could be extracted from the resume.")

    contact = extract_contact_info(text)
    skills = extract_skills(text, skills_list)
    education = extract_education(text)
    experience = extract_experience(text)
    certifications = extract_certifications(text)

    job_match_score = calculate_job_match_score(text, job_description)
    skill_match_score = calculate_skill_match_score(skills, skills_list)
    final_score = round((job_match_score * 0.6) + (skill_match_score * 0.4), 2)

    return {
        "resume_text": text,
        "name": contact["name"],
        "email": contact["email"],
        "phone": contact["phone"],
        "skills": skills,
        "education": education,
        "experience": experience,
        "certifications": certifications,
        "job_match_score": job_match_score,
        "skill_match_score": skill_match_score,
        "final_ranking_score": final_score
    }


# ---------- API Endpoint ----------
@app.post("/rank-resume")
async def rank_resume(
    file: UploadFile = File(...),
    skills: str = Form(...),
    job_description: str = Form(...)
):
    skills_list = [s.strip() for s in skills.split(",") if s.strip()]
    file_stream = io.BytesIO(await file.read())
    file_stream.seek(0)

    try:
        parsed_data = parse_resume(file_stream, file.filename, skills_list, job_description)
        return {"status": "success", "data": parsed_data}
    except Exception as e:
        return {"status": "error", "message": str(e)}
